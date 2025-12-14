from docx import Document

def generate_question_paper(mcqs, filename):
    doc = Document()
    doc.add_heading("MCQ Question Paper", level=1)

    for idx, q in enumerate(mcqs, 1):
        doc.add_paragraph(f"{idx}. {q['question']}")
        for key, value in q["options"].items():
            doc.add_paragraph(f"   {key}. {value}")

    doc.save(filename)

def generate_answer_sheet(mcqs, filename):
    doc = Document()
    doc.add_heading("Answer Sheet", level=1)

    for idx, q in enumerate(mcqs, 1):
        doc.add_paragraph(f"{idx}. {q['answer']}")

    doc.save(filename)
